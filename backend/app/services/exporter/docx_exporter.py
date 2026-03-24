from pathlib import Path

from docx import Document

from app.models.entities import Protocol


def _task_assignees(task) -> str:
    if getattr(task, "assignees_display", None):
        return task.assignees_display
    assignees = getattr(task, "assignees_normalized", None) or []
    if assignees:
        return ", ".join(assignees)
    return task.assignee_b24_name or task.assignee_raw or "-"


def _task_deadline(task) -> str:
    if task.deadline_iso:
        return task.deadline_iso
    if task.deadline_raw:
        return task.deadline_raw
    if task.deadline_note:
        return task.deadline_note
    return "-"


def export_protocol_docx(protocol: Protocol, output_path: Path) -> Path:
    doc = Document()
    title = Path(protocol.original_filename).stem if getattr(protocol, "original_filename", None) else f"Протокол #{protocol.id}"
    created = getattr(protocol, "created_at", None)
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Дата: {created.strftime('%d.%m.%Y') if created else '-'}")

    all_items = sorted(protocol.tasks, key=lambda x: x.order_index)
    agenda_items = [item for item in all_items if getattr(item, "item_kind", "task") in {"agenda", "skipped_agenda"}]
    task_items = [item for item in all_items if getattr(item, "item_kind", "task") == "task"]

    if agenda_items:
        doc.add_heading("Повестка", level=2)
        for idx, agenda in enumerate(agenda_items, start=1):
            status_suffix = " (не обсуждался)" if getattr(agenda, "skipped_discussion_flag", False) else ""
            doc.add_paragraph(f"{idx}. {agenda.normalized_text}{status_suffix}")
    else:
        unique_contexts = []
        seen = set()
        for task in task_items:
            context = (getattr(task, "parent_context", None) or "").strip()
            if context and context not in seen:
                unique_contexts.append(context)
                seen.add(context)
        if unique_contexts:
            doc.add_heading("Повестка", level=2)
            for idx, context in enumerate(unique_contexts, start=1):
                doc.add_paragraph(f"{idx}. {context}")

    discussed_contexts: list[str] = []
    for task in task_items:
        context = (getattr(task, "parent_context", None) or "").strip()
        if context and context not in discussed_contexts:
            discussed_contexts.append(context)
    if discussed_contexts:
        doc.add_heading("Обсудили", level=2)
        for idx, context in enumerate(discussed_contexts, start=1):
            doc.add_paragraph(f"{idx}. {context}")

    doc.add_heading("Поручения", level=2)
    grouped: dict[str, list] = {}
    for task in task_items:
        grouped.setdefault((getattr(task, "parent_context", None) or "").strip() or "Без контекста", []).append(task)

    task_counter = 1
    for context, context_tasks in grouped.items():
        doc.add_paragraph(context).bold = True
        for task in context_tasks:
            doc.add_paragraph(f"{task_counter}. {task.normalized_text}")
            doc.add_paragraph(f"Срок: {_task_deadline(task)}")
            doc.add_paragraph(f"Исполнители: {_task_assignees(task)}")
            if getattr(task, "coordinator", None):
                doc.add_paragraph(f"Координатор: {task.coordinator}")
            task_counter += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
